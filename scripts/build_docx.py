# -*- coding: utf-8 -*-
"""
Build BaoCao_DA1.docx from the HUST template.
- Inherits all styles + numbering (CHUONG/Cap1/Cap2, Caption SEQ, TOC) from the template copy.
- Rebuilds cover + front matter + chapters (parsed from docs/report/*.md) + appendices.
- Captions and TOC/TOF use real Word field codes (replicated verbatim from the template),
  so numbering is automatic and updates on Ctrl+A → F9 in Word.
"""
import sys, io, os, shutil, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE = os.path.join(ROOT, "fileBaoCao", "Mẫu ĐATN_2019_version 1_1 (2).docx")
OUT = os.path.join(ROOT, "fileBaoCao", "BaoCao_DA1.docx")
REPORT = os.path.join(ROOT, "docs", "report")

STUDENT = dict(
    name="NGUYỄN HỮU LINH", mssv="20237455",
    email="linh.nh237455@sis.hust.edu.vn",
    nganh="Công nghệ thông tin",
    chuyennganh="Hệ thống Thông tin Quản lý",
    gvhd="TS. Trần Anh Tú",
    title=("XÂY DỰNG HỆ THỐNG XÁC THỰC HAI YẾU TỐ TÍCH HỢP "
           "QUẢN LÝ MẬT KHẨU KEEPASS VÀ ĐĂNG NHẬP KHÔNG MẬT KHẨU "
           "DỰA TRÊN BẰNG CHỨNG KHÔNG TIẾT LỘ TRI THỨC"),
    doctype="ĐỒ ÁN I", thang="6", nam="2026",
)

# column text width (A4 21cm - left 3.5 - right 2.5 = 15cm)
COL_W = 15.0
FIG_MAX_W = 12.0   # cm, ~ under 75% would be 11.25; we use a readable cap

# ── low-level helpers ────────────────────────────────────────────────────────
def set_run(r, bold=None, size=None, italic=None, color=None, font="Times New Roman"):
    if font: r.font.name = font
    if bold is not None: r.bold = bold
    if italic is not None: r.italic = italic
    if size is not None: r.font.size = Pt(size)
    if color is not None: r.font.color.rgb = RGBColor(*color)
    return r

def add_field(paragraph, instr, cached="", noProof=True):
    """Append a Word field (begin/instr/separate/cached/end) to a paragraph."""
    def mkrun():
        r = OxmlElement('w:r'); paragraph._p.append(r); return r
    r1 = mkrun(); fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'begin'); r1.append(fc)
    r2 = mkrun(); it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr; r2.append(it)
    r3 = mkrun(); fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'separate'); r3.append(fc2)
    r4 = mkrun()
    if noProof:
        rpr = OxmlElement('w:rPr'); np = OxmlElement('w:noProof'); rpr.append(np); r4.append(rpr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = cached; r4.append(t)
    r5 = mkrun(); fc3 = OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'), 'end'); r5.append(fc3)

def set_numpr(paragraph, ilvl, numId):
    """Force multilevel numbering on a heading paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    # remove existing numPr
    for old in pPr.findall(qn('w:numPr')):
        pPr.remove(old)
    numPr = OxmlElement('w:numPr')
    il = OxmlElement('w:ilvl'); il.set(qn('w:val'), str(ilvl)); numPr.append(il)
    ni = OxmlElement('w:numId'); ni.set(qn('w:val'), str(numId)); numPr.append(ni)
    pPr.append(numPr)

def page_break_before(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pbb = OxmlElement('w:pageBreakBefore'); pPr.append(pbb)

def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run(); r.add_break(WD_BREAK.PAGE)

# ── document-level helpers ───────────────────────────────────────────────────
class Builder:
    def __init__(self, doc):
        self.doc = doc
        self.fig_counter = {}   # chapter -> count
        self.tbl_counter = {}
        self.cur_chapter = 0

    def para(self, text="", style=None, align=None, first_indent=True):
        p = self.doc.add_paragraph(style=style)
        if text:
            r = p.add_run(text); set_run(r, size=13)
        if align is not None:
            p.alignment = align
        return p

    def heading(self, text, level, numbered=True, page_break=False):
        style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
        p = self.doc.add_paragraph(style=style)
        r = p.add_run(text); set_run(r, bold=True)
        if numbered:
            set_numpr(p, level - 1, 13)
            if level == 1:
                self.cur_chapter += 1
        else:
            set_numpr(p, 0, 0)  # suppress numbering
        if page_break:
            page_break_before(p)
        return p

    def frontmatter_heading(self, text):
        """Chapter-like heading for front matter, but NOT a Heading style so it
        stays out of the auto Table of Contents."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text.upper()); set_run(r, bold=True, size=14)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.first_line_indent = Cm(0)
        return p

    def bullet(self, text):
        p = self.doc.add_paragraph(style="List Paragraph")
        # use a real bullet via numbering? Template has List Paragraph; add dash bullet
        r = p.add_run("- " + text); set_run(r, size=13)
        p.paragraph_format.left_indent = Cm(0.8)
        return p

    def note(self, text):
        p = self.doc.add_paragraph()
        r = p.add_run(text); set_run(r, size=12, italic=True, color=(0x55,0x55,0x55))
        return p

    def code(self, lines):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(6)
        pPr = p._p.get_or_add_pPr()
        # light shading
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),'F2F2F2'); pPr.append(shd)
        for i, ln in enumerate(lines):
            r = p.add_run(ln)
            set_run(r, size=10.5, font="Consolas", color=(0x1a,0x1a,0x1a))
            if i < len(lines)-1:
                r.add_break()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p

    def figure(self, path, caption):
        full = os.path.normpath(os.path.join(REPORT, path))
        # compute size
        with Image.open(full) as im:
            w, h = im.size
        ratio = h / w
        if ratio > 1.15:   # portrait
            height_cm = 11.0
            width_cm = height_cm / ratio
        else:
            width_cm = FIG_MAX_W
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(full, width=Cm(width_cm))
        # caption below, style Caption
        self._caption(p_after=True, label="Hình", caption=caption)

    def _caption(self, p_after, label, caption):
        ch = self.cur_chapter
        if label == "Hình":
            self.fig_counter[ch] = self.fig_counter.get(ch, 0) + 1
            seq = self.fig_counter[ch]
        else:
            self.tbl_counter[ch] = self.tbl_counter.get(ch, 0) + 1
            seq = self.tbl_counter[ch]
        cap = self.doc.add_paragraph(style="Caption")
        r0 = cap.add_run(label + " "); set_run(r0, size=12)
        add_field(cap, r' STYLEREF 1 \s ', cached=str(ch))
        rdot = cap.add_run("."); set_run(rdot, size=12)
        add_field(cap, r' SEQ %s \* ARABIC \s 1 ' % label, cached=str(seq))
        rsp = cap.add_run(" "); set_run(rsp, size=12)
        rtext = cap.add_run(caption); set_run(rtext, size=12)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return cap

    def table(self, rows, caption):
        # caption ABOVE
        self._caption(p_after=False, label="Bảng", caption=caption)
        ncols = len(rows[0])
        t = self.doc.add_table(rows=len(rows), cols=ncols)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ri, row in enumerate(rows):
            for ci, cell in enumerate(row):
                c = t.cell(ri, ci)
                c.text = ""
                para = c.paragraphs[0]
                # strip markdown bold ** and inline code `
                txt = cell.replace("**", "").replace("`", "")
                r = para.add_run(txt)
                set_run(r, size=12, bold=(ri == 0))
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return t

    def toc(self, instr, placeholder):
        p = self.doc.add_paragraph()
        add_field(p, instr, cached=placeholder)
        return p

# ── markdown parsing ─────────────────────────────────────────────────────────
NUM_PREFIX = re.compile(r'^\d+(\.\d+)*\.?\s+')

def parse_md(path):
    """Return list of block tokens."""
    text = io.open(path, encoding="utf-8").read()
    lines = text.split("\n")
    blocks = []
    i = 0
    while i < len(lines):
        ln = lines[i]
        s = ln.strip()
        if not s:
            i += 1; continue
        # code fence
        if s.startswith("```"):
            j = i + 1; code = []
            while j < len(lines) and not lines[j].strip().startswith("```"):
                code.append(lines[j]); j += 1
            blocks.append(("code", code)); i = j + 1; continue
        # headings
        if s.startswith("### "):
            blocks.append(("h3", strip_num(s[4:]))); i += 1; continue
        if s.startswith("## "):
            blocks.append(("h2", strip_num(s[3:]))); i += 1; continue
        if s.startswith("# "):
            blocks.append(("h1", strip_chapter(s[2:]))); i += 1; continue
        # image  ![cap](path)
        m = re.match(r'!\[(.*?)\]\((.*?)\)', s)
        if m:
            blocks.append(("fig", (m.group(2), m.group(1)))); i += 1; continue
        # table caption marker
        if s.startswith("TBL:"):
            cap = s[4:].strip()
            # collect following table rows
            j = i + 1; rows = []
            while j < len(lines) and lines[j].strip().startswith("|"):
                rows.append(lines[j].strip()); j += 1
            blocks.append(("table", (cap, rows))); i = j; continue
        # standalone table without marker (shouldn't happen) -> treat as table no caption
        if s.startswith("|"):
            j = i; rows = []
            while j < len(lines) and lines[j].strip().startswith("|"):
                rows.append(lines[j].strip()); j += 1
            blocks.append(("table", (None, rows))); i = j; continue
        # blockquote / note
        if s.startswith(">"):
            blocks.append(("note", s.lstrip("> ").strip())); i += 1; continue
        # bullet
        if s.startswith("- "):
            blocks.append(("bullet", s[2:].strip())); i += 1; continue
        # normal paragraph (single line; md paragraphs separated by blank lines)
        blocks.append(("p", s)); i += 1
    return blocks

def strip_num(s):
    return NUM_PREFIX.sub("", s).strip()

def strip_chapter(s):
    s = re.sub(r'^CHƯƠNG\s+\d+\.?\s*', '', s, flags=re.IGNORECASE)
    return s.strip()

def md_table_to_rows(raw_rows):
    rows = []
    for idx, r in enumerate(raw_rows):
        cells = [c.strip() for c in r.strip().strip("|").split("|")]
        if all(set(c) <= set("-: ") for c in cells):  # separator row
            continue
        rows.append(cells)
    return rows

def clean_inline(s):
    # strip markdown emphasis/code markers for plain runs
    s = s.replace("**", "").replace("`", "").replace("*", "")
    return s

def add_bib_runs(p, text):
    """Emit a bibliography paragraph at 13pt, rendering *italic* spans (journal,
    conference and report titles) as real italic runs. No other inline markdown
    is expected in reference entries."""
    for part in re.split(r'(\*[^*]+\*)', text):
        if not part:
            continue
        if len(part) >= 2 and part[0] == '*' and part[-1] == '*':
            set_run(p.add_run(part[1:-1]), size=13, italic=True)
        else:
            set_run(p.add_run(part.replace('`', '')), size=13)

# ── build ────────────────────────────────────────────────────────────────────
def clear_body(doc):
    body = doc.element.body
    for child in list(body.iterchildren()):
        if child.tag == qn('w:sectPr'):
            continue
        body.remove(child)

def build_cover(b):
    d = b.doc
    def cp(text, size, bold=True, sp_after=2):
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); set_run(r, bold=bold, size=size)
        p.paragraph_format.space_after = Pt(sp_after)
        p.paragraph_format.first_line_indent = Cm(0)
        return p
    cp("ĐẠI HỌC BÁCH KHOA HÀ NỘI", 15)
    cp("KHOA TOÁN - TIN", 14)
    # logo placeholder
    logo = os.path.join(ROOT, "docs", "evidence", "fig", "hust-logo.png")
    if os.path.exists(logo):
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(logo, width=Cm(3))
    for _ in range(2): d.add_paragraph()
    cp(STUDENT["doctype"], 24)
    for _ in range(1): d.add_paragraph()
    cp(STUDENT["title"], 17)
    for _ in range(2): d.add_paragraph()
    cp(STUDENT["name"], 14)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Email: " + STUDENT["email"]), size=13, bold=False)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Mã sinh viên: " + STUDENT["mssv"]), size=13, bold=False)
    cp("Chuyên ngành " + STUDENT["chuyennganh"], 13)
    for _ in range(1): d.add_paragraph()
    # GVHD table
    t = d.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.cell(0,0).width = Cm(6); t.cell(0,1).width = Cm(7)
    set_run(t.cell(0,0).paragraphs[0].add_run("Giảng viên hướng dẫn:"), size=13, bold=True)
    set_run(t.cell(0,1).paragraphs[0].add_run(STUDENT["gvhd"]), size=13, bold=True)
    for _ in range(2): d.add_paragraph()
    cp("HÀ NỘI, %s/%s" % (STUDENT["thang"], STUDENT["nam"]), 13)

def build_front_matter(b):
    d = b.doc
    # Nhận xét GVHD
    add_page_break(d)
    b.frontmatter_heading("NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN")
    for _ in range(10): d.add_paragraph()
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(p.add_run("Hà Nội, ngày … tháng … năm 2026"), size=13, italic=True)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(p.add_run("Giảng viên hướng dẫn"), size=13, bold=True)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(p.add_run("(Ký và ghi rõ họ tên)"), size=12, italic=True)

    # Lời cảm ơn + Tóm tắt
    ap = parse_md(os.path.join(REPORT, "appendix.md"))
    add_page_break(d)
    emit_named_section(b, ap, "LỜI CẢM ƠN", "LỜI CẢM ƠN")
    add_page_break(d)
    emit_named_section(b, ap, "TÓM TẮT NỘI DUNG ĐỒ ÁN", "TÓM TẮT NỘI DUNG ĐỒ ÁN")
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(p.add_run("Sinh viên thực hiện"), size=13, bold=True)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(p.add_run("(Ký và ghi rõ họ tên)"), size=12, italic=True)

    # MỤC LỤC
    add_page_break(d)
    b.frontmatter_heading("MỤC LỤC")
    b.toc(r' TOC \o "1-3" \h \z \u ', "Nhấn Ctrl+A rồi F9 để cập nhật mục lục tự động.")
    # DANH MỤC HÌNH VẼ
    add_page_break(d)
    b.frontmatter_heading("DANH MỤC HÌNH VẼ")
    b.toc(r' TOC \h \z \c "Hình" ', "Nhấn Ctrl+A rồi F9 để cập nhật danh mục hình.")
    # DANH MỤC BẢNG BIỂU
    add_page_break(d)
    b.frontmatter_heading("DANH MỤC BẢNG BIỂU")
    b.toc(r' TOC \h \z \c "Bảng" ', "Nhấn Ctrl+A rồi F9 để cập nhật danh mục bảng.")
    # DANH MỤC TỪ VIẾT TẮT
    add_page_break(d)
    b.frontmatter_heading("DANH MỤC TỪ VIẾT TẮT")
    abbr = [
        ["Từ viết tắt", "Giải nghĩa"],
        ["2FA", "Two-Factor Authentication - Xác thực hai yếu tố"],
        ["OTP", "One-Time Password - Mật khẩu dùng một lần"],
        ["HOTP", "HMAC-based One-Time Password (RFC 4226)"],
        ["TOTP", "Time-based One-Time Password (RFC 6238)"],
        ["HMAC", "Hash-based Message Authentication Code"],
        ["AES-GCM", "Advanced Encryption Standard - Galois/Counter Mode"],
        ["JWT", "JSON Web Token"],
        ["ZKP", "Zero-Knowledge Proof - Bằng chứng không tiết lộ tri thức"],
        ["NIZK", "Non-Interactive Zero-Knowledge"],
        ["PWA", "Progressive Web App"],
        ["KDF", "Key Derivation Function - Hàm dẫn xuất khóa"],
        ["API", "Application Programming Interface"],
    ]
    tt = d.add_table(rows=len(abbr), cols=2); tt.style = "Table Grid"; tt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(abbr):
        for ci, cell in enumerate(row):
            tt.cell(ri,ci).text = ""
            r = tt.cell(ri,ci).paragraphs[0].add_run(cell); set_run(r, size=12, bold=(ri==0))

def emit_named_section(b, blocks, heading_text, marker):
    """Emit the paragraphs that appear under a '## marker' heading in appendix.md."""
    b.frontmatter_heading(heading_text)
    capture = False
    for kind, val in blocks:
        if kind == "h2":
            capture = (val.strip().upper() == marker.upper())
            continue
        if kind in ("h1",):
            capture = False; continue
        if not capture:
            continue
        emit_block(b, kind, val)

def emit_block(b, kind, val):
    if kind == "p":
        b.para(clean_inline(val), align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    elif kind == "bullet":
        b.bullet(clean_inline(val))
    elif kind == "note":
        b.note(clean_inline(val))
    elif kind == "code":
        b.code(val)
    elif kind == "h2":
        b.heading(val, 2)
    elif kind == "h3":
        b.heading(val, 3)
    elif kind == "fig":
        path, cap = val
        b.figure(path, cap)
    elif kind == "table":
        cap, raw = val
        rows = md_table_to_rows(raw)
        if rows:
            b.table(rows, cap or "")

def emit_chapter(b, mdfile):
    blocks = parse_md(os.path.join(REPORT, mdfile))
    first_h1 = True
    for kind, val in blocks:
        if kind == "h1":
            b.heading(val, 1, numbered=True, page_break=True)
            first_h1 = False
        else:
            emit_block(b, kind, val)

def emit_appendix_tail(b):
    """Emit Phụ lục A, Phụ lục B, and References from appendix.md."""
    blocks = parse_md(os.path.join(REPORT, "appendix.md"))
    d = b.doc
    # find the sections after front matter: PHỤ LỤC A, PHỤ LỤC B, TÀI LIỆU THAM KHẢO
    mode = None
    for kind, val in blocks:
        if kind == "h2":
            up = val.strip().upper()
            if up.startswith("PHỤ LỤC A"):
                add_page_break(d); b.heading("PHỤ LỤC A - MÃ NGUỒN CHƯƠNG TRÌNH (TRÍCH)", 1, numbered=False)
                mode = "A"; continue
            if up.startswith("PHỤ LỤC B"):
                add_page_break(d); b.heading("PHỤ LỤC B - HƯỚNG DẪN CÀI ĐẶT VÀ SỬ DỤNG", 1, numbered=False)
                mode = "B"; continue
            if up.startswith("TÀI LIỆU THAM KHẢO"):
                add_page_break(d); b.heading("TÀI LIỆU THAM KHẢO", 1, numbered=False)
                mode = "REF"; continue
            if up in ("TÓM TẮT NỘI DUNG ĐỒ ÁN", "LỜI CẢM ƠN"):
                mode = None; continue
            if mode in ("A", "B"):
                b.heading(val, 2, numbered=False)  # text already carries A.1/B.1
                continue
        if mode is None:
            continue
        if mode == "REF":
            if kind == "p":
                p = b.doc.add_paragraph(style="Bibliography")
                add_bib_runs(p, val)
                # Left-align (not justify) so long URLs in [9]/[12] don't blow up
                # inter-word spacing on the preceding line.
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif kind == "note":
                b.note(clean_inline(val))
            continue
        if mode in ("A", "B") and kind == "h3":
            b.heading(val, 3, numbered=False)  # text already carries A.x/B.x label
            continue
        emit_block(b, kind, val)

def enable_update_fields(doc):
    settings = doc.settings.element
    uf = settings.find(qn('w:updateFields'))
    if uf is None:
        uf = OxmlElement('w:updateFields'); uf.set(qn('w:val'), 'true')
        settings.insert(0, uf)

def _clear_footer(footer):
    """Remove ALL content from the footer story, including the template's
    page-number content control (w:sdt) which footer.paragraphs cannot see."""
    footer.is_linked_to_previous = False
    ftr = footer._element
    for child in list(ftr):
        ftr.remove(child)

def add_footer_pagenum(doc):
    """Single centered page number; cover page (first page) has none.
    The template carries its own page number inside a w:sdt content control,
    so we wipe the footer story first to avoid a duplicate."""
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    # main footer (pages 2+): wipe, then one centered PAGE field
    _clear_footer(sec.footer)
    p = sec.footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(p, ' PAGE ', cached="1")
    # first page (cover) footer: wipe and leave a single empty paragraph
    _clear_footer(sec.first_page_footer)
    sec.first_page_footer.add_paragraph()

def main():
    shutil.copyfile(TEMPLATE, OUT)
    doc = Document(OUT)
    clear_body(doc)
    b = Builder(doc)
    build_cover(b)
    build_front_matter(b)
    for f in ["chuong-1.md", "chuong-2.md", "chuong-3.md", "chuong-4.md", "chuong-5.md"]:
        emit_chapter(b, f)
    emit_appendix_tail(b)
    add_footer_pagenum(doc)
    enable_update_fields(doc)
    doc.save(OUT)
    # summary
    print("Saved:", OUT)
    print("chapters:", b.cur_chapter, "| figures/chap:", b.fig_counter, "| tables/chap:", b.tbl_counter)

if __name__ == "__main__":
    main()
